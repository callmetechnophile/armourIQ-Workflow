import os
import csv
import uuid
import json
from datetime import datetime
from typing import Dict, Any, List

from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen.canvas import Canvas

# ---------------------------------------------------------------------------
# Font Registration
# ---------------------------------------------------------------------------
FONTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "fonts")

CODE_FONT = "Courier"
CODE_FONT_BOLD = "Courier-Bold"

try:
    ubuntu_regular = os.path.join(FONTS_DIR, "Ubuntu-Regular.ttf")
    ubuntu_bold = os.path.join(FONTS_DIR, "Ubuntu-Bold.ttf")
    if os.path.exists(ubuntu_regular):
        pdfmetrics.registerFont(TTFont("Ubuntu", ubuntu_regular))
        CODE_FONT = "Ubuntu"
    if os.path.exists(ubuntu_bold):
        pdfmetrics.registerFont(TTFont("Ubuntu-Bold", ubuntu_bold))
        CODE_FONT_BOLD = "Ubuntu-Bold"
except Exception as e:
    print(f"[Font] Ubuntu registration failed, using Courier fallback: {e}")

EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# Numbered Canvas – renders "Page X of Y" footer on every page
# ---------------------------------------------------------------------------
_footer_project_title = ""

class NumberedCanvas(Canvas):
    def __init__(self, *args, **kwargs):
        Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer(num_pages)
            Canvas.showPage(self)
        Canvas.save(self)

    def _draw_footer(self, page_count):
        self.saveState()
        self.setFont("Times-Roman", 9)
        self.setFillColor(colors.HexColor("#475569"))
        self.setStrokeColor(colors.HexColor("#d1d5db"))
        self.setLineWidth(0.4)
        self.line(40, 42, A4[0] - 40, 42)
        self.drawString(40, 28, f"{_footer_project_title} — Engineering Report")
        self.drawRightString(A4[0] - 40, 28, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()


# ---------------------------------------------------------------------------
# generate_project_title  (Gemma 2 → local fallback)
# ---------------------------------------------------------------------------
def generate_project_title(user_query: str) -> str:
    groq_api_key = os.environ.get("GROQ_API_KEY")
    if groq_api_key and not groq_api_key.startswith("gsk_placeholder"):
        try:
            import httpx
            payload = {
                "model": "gemma2-9b-it",
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You are a technical naming system. Convert the user's hardware prototype query "
                            "into a short, clean, professional engineering-style project title. "
                            "Rules: Remove filler words, remove 'I want to build', remove generic words. "
                            "Compress into 1–3 words. Must sound technical (e.g. 'SolarVac System', 'AeroDrop System'). "
                            "Reply ONLY with the project name. No quotes, periods, explanation or markdown."
                        )
                    },
                    {"role": "user", "content": user_query}
                ],
                "temperature": 0.1,
                "max_tokens": 10
            }
            headers = {"Authorization": f"Bearer {groq_api_key}", "Content-Type": "application/json"}
            with httpx.Client(timeout=10.0) as client:
                res = client.post("https://api.groq.com/openai/v1/chat/completions", json=payload, headers=headers)
                if res.status_code == 200:
                    title = res.json()["choices"][0]["message"]["content"].strip().replace('"', '').replace("'", "").replace('.', '').strip()
                    if title:
                        return title
        except Exception:
            pass

    q = user_query.lower()
    if "solar" in q and ("vacuum" in q or "clean" in q):
        return "SolarVac System"
    if "drone" in q or "delivery" in q:
        return "AeroDrop System"
    if "irrigation" in q or "water" in q or "agri" in q:
        return "AgriFlow System"
    if "hand" in q or "bionic" in q or "robotic" in q:
        return "BionicHand System"

    words = [w for w in user_query.split() if w.lower() not in {"i", "want", "to", "build", "a", "an", "the", "using", "with", "for", "smart"}]
    if len(words) >= 2:
        return "".join(w.capitalize() for w in words[:2]) + " System"
    if words:
        return words[0].capitalize() + " System"
    return "ArmourFlow Project"


# ---------------------------------------------------------------------------
# generate_abstract  (Nemotron → local fallback)
# ---------------------------------------------------------------------------
def generate_abstract(project_title: str, user_query: str) -> str:
    groq_api_key = os.environ.get("GROQ_API_KEY")
    if groq_api_key and not groq_api_key.startswith("gsk_placeholder"):
        try:
            import httpx
            payload = {
                "model": "llama-3.1-nemotron-70b-specdec",
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "You are a professional hardware research engineer. "
                            f"Write a project abstract of 150–250 words for '{project_title}' (Query: '{user_query}'). "
                            "Include four sections: Problem, Solution, Implementation Idea, Feasibility. "
                            "Keep it technical, objective, and clear. No markdown headers or greetings."
                        )
                    }
                ],
                "temperature": 0.2,
                "max_tokens": 400
            }
            headers = {"Authorization": f"Bearer {groq_api_key}", "Content-Type": "application/json"}
            with httpx.Client(timeout=12.0) as client:
                res = client.post("https://api.groq.com/openai/v1/chat/completions", json=payload, headers=headers)
                if res.status_code == 200:
                    abstract = res.json()["choices"][0]["message"]["content"].strip()
                    if abstract:
                        return abstract
        except Exception:
            pass

    return (
        f"PROBLEM: Standard prototyping pipelines for {user_query} suffer from fragmented component sourcing, "
        "voltage-domain mismatches, and insufficient compliance auditing during autonomous research phases.\n\n"
        f"SOLUTION: {project_title} resolves these bottlenecks via an autonomous multi-agent pipeline coupled with "
        "the ArmorIQ zero-trust enforcer, dynamically optimising vendor logistics and validating pin configurations.\n\n"
        "IMPLEMENTATION: The physical system uses core microcontroller logic, PCA9685 servo drivers, and "
        "high-efficiency power cells mapped via custom I2C/SPI wire harnesses.\n\n"
        "FEASIBILITY: Engineering validation estimates readiness at 85 % with negligible voltage-mismatch hazards."
    )


# ---------------------------------------------------------------------------
# generate_firmware_code
# ---------------------------------------------------------------------------
def generate_firmware_code(components: List[Dict[str, Any]], intent: str) -> str:
    has_pca = any("pca9685" in (c.get("component") or c.get("name", "")).lower() for c in components)
    has_servos = any("servo" in (c.get("component") or c.get("name", "")).lower() for c in components)

    if has_pca and has_servos:
        return """#include <Wire.h>
#include <Adafruit_PWMServoDriver.h>

Adafruit_PWMServoDriver pwm = Adafruit_PWMServoDriver();

#define FLEX_THUMB  32
#define FLEX_INDEX  33
#define FLEX_MIDDLE 34
#define FLEX_RING   35
#define FLEX_PINKY  36
#define SERVOMIN  150
#define SERVOMAX  600

void setup() {
  Serial.begin(115200);
  Wire.begin(21, 22);
  pwm.begin();
  pwm.setOscillatorFrequency(27000000);
  pwm.setPWMFreq(50);
  pinMode(FLEX_THUMB, INPUT);
  pinMode(FLEX_INDEX, INPUT);
  pinMode(FLEX_MIDDLE, INPUT);
  pinMode(FLEX_RING, INPUT);
  pinMode(FLEX_PINKY, INPUT);
}

void loop() {
  int vals[] = {
    analogRead(FLEX_THUMB),  analogRead(FLEX_INDEX),
    analogRead(FLEX_MIDDLE), analogRead(FLEX_RING),
    analogRead(FLEX_PINKY)
  };
  for (int i = 0; i < 5; i++) {
    int pulse = map(vals[i], 1000, 3000, SERVOMIN, SERVOMAX);
    pwm.setPWM(i, 0, constrain(pulse, SERVOMIN, SERVOMAX));
  }
  delay(20);
}"""
    if "solar" in intent.lower() or "vacuum" in intent.lower():
        return """#define VACUUM_MOTOR_PWM  15
#define BATTERY_ADC_PIN   35

void setup() {
  Serial.begin(115200);
  pinMode(VACUUM_MOTOR_PWM, OUTPUT);
}

void loop() {
  float voltage = (analogRead(BATTERY_ADC_PIN) / 4095.0) * 3.3 * 4.0;
  analogWrite(VACUUM_MOTOR_PWM, voltage < 10.5 ? 0 : 180);
  delay(100);
}"""
    return """#define STATUS_LED 2

void setup() {
  Serial.begin(115200);
  pinMode(STATUS_LED, OUTPUT);
}

void loop() {
  digitalWrite(STATUS_LED, HIGH); delay(500);
  digitalWrite(STATUS_LED, LOW);  delay(500);
}"""


# ===================================================================
# MAIN PDF BUILDER
# ===================================================================
def export_pdf(data: Dict[str, Any]) -> Dict[str, Any]:
    global _footer_project_title

    intent = data.get("intent", "Project Specification")
    project_title = generate_project_title(intent)
    _footer_project_title = project_title
    abstract_text = generate_abstract(project_title, intent)

    filename = f"{project_title.replace(' ', '_')}.pdf"
    file_path = os.path.join(EXPORT_DIR, filename)

    doc = SimpleDocTemplate(file_path, pagesize=A4,
                            rightMargin=40, leftMargin=40,
                            topMargin=40, bottomMargin=55)
    story: list = []
    styles = getSampleStyleSheet()

    # ------------------------------------------------------------------
    # Typography — Times New Roman 12, justified
    # ------------------------------------------------------------------
    title_style = ParagraphStyle("DocTitle", parent=styles["Heading1"],
        fontName="Times-Bold", fontSize=22, leading=26,
        textColor=colors.HexColor("#0f172a"), alignment=TA_CENTER, spaceAfter=10)

    subtitle_style = ParagraphStyle("DocSub", parent=styles["Heading2"],
        fontName="Times-Bold", fontSize=14, leading=18,
        textColor=colors.HexColor("#475569"), alignment=TA_CENTER, spaceAfter=8)

    tagline_style = ParagraphStyle("DocTag", parent=styles["Normal"],
        fontName="Times-Italic", fontSize=12, leading=16,
        textColor=colors.HexColor("#64748b"), alignment=TA_CENTER, spaceAfter=6)

    section_style = ParagraphStyle("Section", parent=styles["Heading2"],
        fontName="Times-Bold", fontSize=14, leading=18,
        textColor=colors.HexColor("#1e293b"), spaceBefore=14, spaceAfter=8,
        alignment=TA_LEFT)

    subsection_style = ParagraphStyle("SubSection", parent=styles["Heading3"],
        fontName="Times-Bold", fontSize=12, leading=16,
        textColor=colors.HexColor("#334155"), spaceBefore=8, spaceAfter=4,
        alignment=TA_LEFT)

    body = ParagraphStyle("Body", parent=styles["Normal"],
        fontName="Times-Roman", fontSize=12, leading=16,
        textColor=colors.HexColor("#1e293b"), spaceAfter=6,
        alignment=TA_JUSTIFY)

    code_style = ParagraphStyle("Code", parent=styles["Normal"],
        fontName=CODE_FONT, fontSize=9, leading=12,
        textColor=colors.HexColor("#0f172a"), spaceAfter=3,
        alignment=TA_LEFT)

    th = ParagraphStyle("TH", parent=styles["Normal"],
        fontName="Times-Bold", fontSize=10, leading=13, textColor=colors.white)
    td = ParagraphStyle("TD", parent=styles["Normal"],
        fontName="Times-Roman", fontSize=10, leading=13,
        textColor=colors.HexColor("#1e293b"))

    # Helper -----------------------------------------------------------
    components = data.get("components", [])
    validation = data.get("validation", {})
    readiness_score = validation.get("readiness_score", 80)
    complexity = "High" if len(components) > 5 else "Medium"
    cost_summary = data.get("cost_summary", {})
    total_cost = cost_summary.get("total_landed_cost", sum(c.get("cost", 0) for c in components))
    if isinstance(total_cost, (int, float)) and total_cost < 1000:
        total_cost = int(total_cost * 83)
    else:
        total_cost = int(total_cost)

    def _tbl(rows, widths):
        t = Table(rows, colWidths=widths)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f8fafc")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8fafc"), colors.white]),
        ]))
        return t

    # ==================================================================
    # PAGE 1 — COVER
    # ==================================================================
    story.append(Spacer(1, 120))
    story.append(Paragraph(project_title, title_style))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Engineering Research &amp; Execution Report", subtitle_style))
    story.append(Spacer(1, 30))
    story.append(Paragraph("Generated via WorkflowGuide AI", tagline_style))
    story.append(Spacer(1, 100))
    story.append(Paragraph(f"<b>Generation Timestamp:</b>  {datetime.utcnow().strftime('%Y-%m-%d  %H:%M UTC')}", body))
    story.append(Paragraph(f"<b>Execution Readiness:</b>  {readiness_score} %", body))
    story.append(Paragraph(f"<b>Total BOM Cost:</b>  ₹{total_cost:,}", body))
    story.append(Paragraph(f"<b>Project Complexity:</b>  {complexity}", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 1 — ABSTRACT / PROBLEM / SOLUTION  (flows naturally)
    # ==================================================================
    story.append(Paragraph("1.  Abstract", section_style))
    story.append(Paragraph(abstract_text, body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("2.  Problem Statement", section_style))
    story.append(Paragraph(f"<b>Target Objective:</b> {intent}", body))
    story.append(Paragraph(
        "Component acquisition is bottlenecked by fragmented distributor channels leading to unoptimised "
        "shipping expenses. Microcontroller-driven 3.3 V CMOS logic connected directly to 5 V TTL motor "
        "controllers causes signal degradation. Autonomous research modules lack real-time sandboxing, "
        "allowing unauthorised tool invocations.", body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("3.  Proposed Solution", section_style))
    story.append(Paragraph(
        f"<b>{project_title}</b> is implemented as a cryptographically-secured control platform. "
        "An integrated planner orchestrates retrieval, extraction, and validation agents while the cost "
        "engine dynamically swaps vendors for optimal landed prices. ArmorIQ intercepts out-of-scope "
        "system calls at runtime, logging tamper-proof cryptographic receipts.", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 2 — SYSTEM ARCHITECTURE + BOM
    # ==================================================================
    story.append(Paragraph("4.  System Architecture", section_style))
    flow = [
        "User Query → Planner Agent → Retrieval Agent → Extraction Agent",
        "→ Procurement Agent → Research Agent → Validation Agent",
        "→ Optimisation Agent → Planning Agent → Export Agent",
        "[ArmorIQ trust enforcer validates every tool call]"
    ]
    for line in flow:
        story.append(Paragraph(line, code_style))
    story.append(Spacer(1, 14))

    story.append(Paragraph("5.  Bill of Materials", section_style))
    bom_rows = [[Paragraph("Component", th), Paragraph("Vendor", th),
                  Paragraph("Base ₹", th), Paragraph("Ship ₹", th),
                  Paragraph("Final ₹", th), Paragraph("Stock", th)]]
    for c in components:
        bom_rows.append([
            Paragraph(c.get("component") or c.get("name", "—"), td),
            Paragraph(c.get("selected_vendor", "element14 India"), td),
            Paragraph(f"₹{int(c.get('base_cost', c.get('cost', 0)*83)):,}", td),
            Paragraph(f"₹{int(c.get('shipping_cost', 90)):,}", td),
            Paragraph(f"₹{int(c.get('final_cost', c.get('cost', 0)*83+90)):,}", td),
            Paragraph(c.get("stock", "In Stock"), td),
        ])
    story.append(_tbl(bom_rows, [150, 85, 55, 50, 55, 55]))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"<b>Grand Total:</b>  ₹{total_cost:,}", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 3 — ALTERNATIVES + POWER
    # ==================================================================
    story.append(Paragraph("6.  Alternative Components", section_style))
    alternatives = data.get("alternatives", [])
    if not alternatives:
        story.append(Paragraph("All components currently at optimum selections.", body))
    else:
        for idx, alt in enumerate(alternatives[:4]):
            story.append(Paragraph(
                f"<b>{idx+1}.</b>  {alt.get('original', '—')}  →  "
                f"{alt.get('alternative') or alt.get('name', '—')}  "
                f"<i>({alt.get('reason', '')})</i>", body))
    story.append(Spacer(1, 14))

    story.append(Paragraph("7.  Power Analysis", section_style))
    power_data = data.get("power_analysis", {})
    power_items = power_data.get("power_items", [])
    p_rows = [[Paragraph("Component", th), Paragraph("Voltage", th),
                Paragraph("Nominal mA", th), Paragraph("Peak mA", th)]]
    for item in power_items:
        p_rows.append([
            Paragraph(item.get("component", ""), td),
            Paragraph(f"{item.get('voltage')} V", td),
            Paragraph(str(item.get("nominal_current", "")), td),
            Paragraph(str(item.get("peak_current", "")), td),
        ])
    story.append(_tbl(p_rows, [200, 80, 90, 90]))
    p_sum = power_data.get("summary", {})
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"Total load: {p_sum.get('total_power_load_w', 1.5)} W  |  "
                           f"Battery runtime: {p_sum.get('estimated_runtime_hours', 2.0)} hrs", body))
    for w in power_data.get("warnings", []):
        story.append(Paragraph(f"⚠  {w}", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 4 — DEPENDENCY + WIRING + PIN TABLE
    # ==================================================================
    story.append(Paragraph("8.  Dependency Graph", section_style))
    dep_edges = data.get("dependency_graph", {}).get("edges", [])
    for edge in dep_edges:
        story.append(Paragraph(
            f"{edge.get('source')}  →  {edge.get('target')}  ({edge.get('type')}: {edge.get('label')})", body))
    if not dep_edges:
        story.append(Paragraph("No explicit dependency links extracted.", body))
    story.append(Spacer(1, 14))

    story.append(Paragraph("9.  Wiring Diagram &amp; Pin Configuration", section_style))
    connections = data.get("wiring_diagram", {}).get("connections", [])
    pin_rows = [[Paragraph("Source", th), Paragraph("Pin", th),
                 Paragraph("Target", th), Paragraph("Pin", th),
                 Paragraph("Protocol", th)]]
    for conn in connections:
        pin_rows.append([
            Paragraph(conn.get("source", ""), td),
            Paragraph(conn.get("source_pin", ""), td),
            Paragraph(conn.get("target", ""), td),
            Paragraph(conn.get("target_pin", ""), td),
            Paragraph(conn.get("protocol", ""), td),
        ])
    story.append(_tbl(pin_rows, [100, 80, 100, 80, 80]))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 5 — DATASHEETS + RESEARCH PAPERS
    # ==================================================================
    story.append(Paragraph("10.  Datasheets", section_style))
    ds_rows = [[Paragraph("Component", th), Paragraph("Datasheet URL", th), Paragraph("Source", th)]]
    for ds in data.get("datasheets", []):
        url = ds.get("datasheet_link", "")
        link = f'<a href="{url}" color="blue"><u>{url[:50]}</u></a>'
        ds_rows.append([
            Paragraph(ds.get("component", ""), td),
            Paragraph(link, td),
            Paragraph(ds.get("source", ""), td),
        ])
    story.append(_tbl(ds_rows, [140, 250, 110]))
    story.append(Spacer(1, 14))

    story.append(Paragraph("11.  Research Papers (Top 3)", section_style))
    papers = data.get("papers", [])
    rp_rows = [[Paragraph("Title", th), Paragraph("Score", th),
                 Paragraph("Year", th), Paragraph("Link", th)]]
    for p in papers[:3]:
        url = p.get("url", "")
        link = f'<a href="{url}" color="blue"><u>Source</u></a>'
        rp_rows.append([
            Paragraph(p.get("title", "—"), td),
            Paragraph(f"{p.get('score', 80)}/100", td),
            Paragraph(str(p.get("publish_year", 2022)), td),
            Paragraph(link, td),
        ])
    story.append(_tbl(rp_rows, [260, 50, 50, 80]))
    paper_summary = data.get("paper_summary", {})
    story.append(Spacer(1, 6))
    story.append(Paragraph(paper_summary.get("summary",
        "Consolidated literature compiled in research logs."), body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 6 — VALIDATION + OPTIMISATION
    # ==================================================================
    story.append(Paragraph("12.  Engineering Validation", section_style))
    for chk in validation.get("validation_checks", []):
        story.append(Paragraph(
            f"<b>[{chk.get('status')}]</b>  {chk.get('check')}:  {chk.get('details')}", body))
    for c in validation.get("contradictions", []):
        story.append(Paragraph(
            f"⚠  <b>{c.get('conflict')} ({c.get('severity')}):</b>  {c.get('explanation')}  "
            f"<i>Mitigation: {c.get('mitigation')}</i>", body))
    story.append(Spacer(1, 14))

    story.append(Paragraph("13.  Optimisation Report", section_style))
    opt = data.get("optimization", {})
    for r in opt.get("recommendations", []):
        story.append(Paragraph(f"•  {r}", body))
    story.append(Paragraph(
        f"Optimisation score: {opt.get('optimization_score', 90)} %  |  "
        f"Cost saved: ₹{int(opt.get('saved_amount', 12)*83):,}", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 7 — GENERATED CODE
    # ==================================================================
    story.append(Paragraph("14.  Generated Controller Firmware", section_style))
    story.append(Paragraph(f"<b>Target MCU:</b> ESP32  |  <b>Framework:</b> Arduino C++", body))
    story.append(Spacer(1, 6))
    code_text = generate_firmware_code(components, intent)
    for line in code_text.split("\n"):
        story.append(Paragraph(line.replace(" ", "&nbsp;").replace("<", "&lt;").replace(">", "&gt;"), code_style))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 8 — GANTT + CONNECTION NOTES
    # ==================================================================
    story.append(Paragraph("15.  Gantt Timeline", section_style))
    for phase in data.get("roadmap", []):
        story.append(Paragraph(
            f"<b>Phase {phase.get('phase')}: {phase.get('title')}</b>  "
            f"({phase.get('duration_days')} days)  —  {phase.get('deliverable', '')}", body))
        story.append(Paragraph(phase.get("description", ""), body))
    story.append(Spacer(1, 14))

    story.append(Paragraph("16.  Connection Assistant Notes", section_style))
    story.append(Paragraph(
        "<b>1. Ground Loop Prevention:</b>  PCA9685 logic ground must share a common node with ESP32 GND; "
        "servo return current should route directly to the LiPo terminal block.", body))
    story.append(Paragraph(
        "<b>2. I²C Pull-ups:</b>  Place 4.7 kΩ pull-up resistors on SDA/SCL to the 3.3 V rail to prevent "
        "bus crashes from servo transient noise.", body))
    story.append(Paragraph(
        "<b>3. Power Decoupling:</b>  Install a 1000 µF low-ESR electrolytic across PCA9685 V+ terminal "
        "to filter ripple during peak servo stall draws.", body))
    story.append(PageBreak())

    # ==================================================================
    # SECTION 9 — AUDIT TRAIL + CONCLUSION
    # ==================================================================
    story.append(Paragraph("17.  ArmorIQ Audit Trail", section_style))
    a_rows = [[Paragraph("Agent", th), Paragraph("Tool", th), Paragraph("Scope", th),
                Paragraph("Status", th), Paragraph("Receipt Hash", th)]]
    for log in data.get("audit_trail", [])[:14]:
        h = log.get("receipt_hash", "—")
        short = f"{h[:6]}…{h[-6:]}" if len(h) > 12 else h
        a_rows.append([
            Paragraph(log.get("agent", ""), td),
            Paragraph(log.get("tool", ""), td),
            Paragraph(log.get("scope", ""), td),
            Paragraph(log.get("status", ""), td),
            Paragraph(short, td),
        ])
    story.append(_tbl(a_rows, [95, 100, 100, 60, 105]))
    story.append(Spacer(1, 14))

    story.append(Paragraph("18.  Conclusion &amp; Execution Readiness", section_style))
    story.append(Paragraph(f"<b>Build Feasibility:</b>  {readiness_score} %", body))
    story.append(Paragraph(f"<b>Cost Feasibility:</b>  92 %", body))
    story.append(Paragraph(f"<b>Complexity:</b>  {complexity}", body))
    story.append(Paragraph("<b>Estimated Build Time:</b>  12 days  (4-phase roadmap)", body))
    story.append(Paragraph("<b>Skill Level:</b>  Intermediate Electronics &amp; Firmware", body))
    story.append(Spacer(1, 10))
    story.append(Paragraph(
        f"The {project_title} project has been compiled and verified by the multi-agent system. "
        "Electrical wiring schematics, I²C logic domains, and battery discharge thresholds show full "
        "compatibility. No critical hazards were logged in the validation audit. Sourcing reflects the "
        "most cost-effective alternatives, making the project fully ready for hardware prototyping.", body))

    # ------------------------------------------------------------------
    # Build PDF with numbered canvas
    # ------------------------------------------------------------------
    doc.build(story, canvasmaker=NumberedCanvas)

    return {"filename": filename, "url": f"/api/exports/{filename}", "status": "SUCCESS"}


# ===================================================================
# CSV / Markdown exporters  (unchanged logic, improved columns)
# ===================================================================
def export_csv(data: Dict[str, Any]) -> Dict[str, Any]:
    fid = f"package_{uuid.uuid4().hex[:8]}"
    fn = f"{fid}.csv"
    fp = os.path.join(EXPORT_DIR, fn)
    with open(fp, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Category", "Component", "Base Cost", "Shipping", "Final Cost", "Vendor", "Notes"])
        for c in data.get("components", []):
            w.writerow([
                c.get("category", ""), c.get("component") or c.get("name", ""),
                c.get("base_cost", c.get("cost", 0)), c.get("shipping_cost", 90),
                c.get("final_cost", c.get("cost", 0) + 90),
                c.get("selected_vendor", "element14 India"), c.get("notes", ""),
            ])
    return {"filename": fn, "url": f"/api/exports/{fn}", "status": "SUCCESS"}


def export_markdown(data: Dict[str, Any]) -> Dict[str, Any]:
    fid = f"package_{uuid.uuid4().hex[:8]}"
    fn = f"{fid}.md"
    fp = os.path.join(EXPORT_DIR, fn)
    lines = [
        f"# WorkflowGuide AI — {data.get('intent', 'Project')}\n",
        f"**Readiness:** {data.get('validation', {}).get('readiness_score', 0)} %\n",
        f"**Risk:** {data.get('validation', {}).get('risk_score', 0)} %\n",
        "\n## Bill of Materials\n",
        "| Category | Component | Base ₹ | Shipping ₹ | Final ₹ | Vendor | Notes |",
        "| --- | --- | --- | --- | --- | --- | --- |",
    ]
    for c in data.get("components", []):
        lines.append(
            f"| {c.get('category')} | {c.get('component') or c.get('name')} | "
            f"₹{c.get('base_cost', c.get('cost', 0)):,} | ₹{c.get('shipping_cost', 90):,} | "
            f"₹{c.get('final_cost', c.get('cost', 0)+90):,} | {c.get('selected_vendor')} | {c.get('notes')} |"
        )
    with open(fp, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return {"filename": fn, "url": f"/api/exports/{fn}", "status": "SUCCESS"}


# ===================================================================
# DOCX EXPORTER
# ===================================================================
def export_docx(data: Dict[str, Any]) -> Dict[str, Any]:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    intent = data.get("intent", "Project Specification")
    project_title = generate_project_title(intent)
    abstract_text = generate_abstract(project_title, intent)

    filename = f"{project_title.replace(' ', '_')}.docx"
    file_path = os.path.join(EXPORT_DIR, filename)

    doc = Document()

    # --- Page setup: A4 ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # --- Style helpers ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Register Ubuntu font for code paragraphs (fallback to Courier New)
    ubuntu_path = os.path.join(FONTS_DIR, "Ubuntu-Regular.ttf")
    code_font_name = "Ubuntu" if os.path.exists(ubuntu_path) else "Courier New"

    def add_heading_styled(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
        return h

    def add_body(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.name = code_font_name
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)
        return p

    def add_table_from_rows(headers, rows_data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        for i, h_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Dark background
            shading = cell._element.get_or_add_tcPr()
            bg = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "1e293b"
            })
            shading.append(bg)
        # Data rows
        for row_data in rows_data:
            row = table.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(val))
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
        return table

    # Context
    components = data.get("components", [])
    validation = data.get("validation", {})
    readiness_score = validation.get("readiness_score", 80)
    complexity = "High" if len(components) > 5 else "Medium"
    cost_summary = data.get("cost_summary", {})
    total_cost = cost_summary.get("total_landed_cost", sum(c.get("cost", 0) for c in components))
    if isinstance(total_cost, (int, float)) and total_cost < 1000:
        total_cost = int(total_cost * 83)
    else:
        total_cost = int(total_cost)

    # ===================== COVER PAGE =====================
    doc.add_paragraph()  # spacing
    doc.add_paragraph()
    h = doc.add_heading(project_title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Engineering Research & Execution Report")
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run("Generated via WorkflowGuide AI")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph()
    add_body(f"Generation Timestamp:  {datetime.utcnow().strftime('%Y-%m-%d  %H:%M UTC')}", bold=True)
    add_body(f"Execution Readiness:  {readiness_score} %", bold=True)
    add_body(f"Total BOM Cost:  ₹{total_cost:,}", bold=True)
    add_body(f"Project Complexity:  {complexity}", bold=True)
    doc.add_page_break()

    # ===================== ABSTRACT + PROBLEM + SOLUTION =====================
    add_heading_styled("1.  Abstract", level=1)
    add_body(abstract_text)

    add_heading_styled("2.  Problem Statement", level=1)
    add_body(f"Target Objective: {intent}", bold=True)
    add_body(
        "Component acquisition is bottlenecked by fragmented distributor channels leading to unoptimised "
        "shipping expenses. Microcontroller-driven 3.3 V CMOS logic connected directly to 5 V TTL motor "
        "controllers causes signal degradation. Autonomous research modules lack real-time sandboxing, "
        "allowing unauthorised tool invocations."
    )

    add_heading_styled("3.  Proposed Solution", level=1)
    add_body(
        f"{project_title} is implemented as a cryptographically-secured control platform. "
        "An integrated planner orchestrates retrieval, extraction, and validation agents while the cost "
        "engine dynamically swaps vendors for optimal landed prices. ArmorIQ intercepts out-of-scope "
        "system calls at runtime, logging tamper-proof cryptographic receipts."
    )
    doc.add_page_break()

    # ===================== ARCHITECTURE + BOM =====================
    add_heading_styled("4.  System Architecture", level=1)
    for line in [
        "User Query → Planner Agent → Retrieval Agent → Extraction Agent",
        "→ Procurement Agent → Research Agent → Validation Agent",
        "→ Optimisation Agent → Planning Agent → Export Agent",
        "[ArmorIQ trust enforcer validates every tool call]"
    ]:
        add_code(line)

    add_heading_styled("5.  Bill of Materials", level=1)
    bom_headers = ["Component", "Vendor", "Base ₹", "Ship ₹", "Final ₹", "Stock"]
    bom_rows = []
    for c in components:
        bom_rows.append([
            c.get("component") or c.get("name", "—"),
            c.get("selected_vendor", "element14 India"),
            f"₹{int(c.get('base_cost', c.get('cost', 0)*83)):,}",
            f"₹{int(c.get('shipping_cost', 90)):,}",
            f"₹{int(c.get('final_cost', c.get('cost', 0)*83+90)):,}",
            c.get("stock", "In Stock"),
        ])
    add_table_from_rows(bom_headers, bom_rows)
    add_body(f"Grand Total:  ₹{total_cost:,}", bold=True)
    doc.add_page_break()

    # ===================== ALTERNATIVES + POWER =====================
    add_heading_styled("6.  Alternative Components", level=1)
    alternatives = data.get("alternatives", [])
    if not alternatives:
        add_body("All components currently at optimum selections.")
    else:
        for idx, alt in enumerate(alternatives[:4]):
            add_body(
                f"{idx+1}.  {alt.get('original', '—')}  →  "
                f"{alt.get('alternative') or alt.get('name', '—')}  "
                f"({alt.get('reason', '')})"
            )

    add_heading_styled("7.  Power Analysis", level=1)
    power_data = data.get("power_analysis", {})
    power_items = power_data.get("power_items", [])
    p_headers = ["Component", "Voltage", "Nominal mA", "Peak mA"]
    p_rows = [[item.get("component", ""), f"{item.get('voltage')} V",
                str(item.get("nominal_current", "")), str(item.get("peak_current", ""))]
               for item in power_items]
    add_table_from_rows(p_headers, p_rows)
    p_sum = power_data.get("summary", {})
    add_body(f"Total load: {p_sum.get('total_power_load_w', 1.5)} W  |  "
             f"Battery runtime: {p_sum.get('estimated_runtime_hours', 2.0)} hrs")
    for w in power_data.get("warnings", []):
        add_body(f"⚠  {w}")
    doc.add_page_break()

    # ===================== DEPENDENCY + WIRING + PINS =====================
    add_heading_styled("8.  Dependency Graph", level=1)
    dep_edges = data.get("dependency_graph", {}).get("edges", [])
    for edge in dep_edges:
        add_body(f"{edge.get('source')}  →  {edge.get('target')}  ({edge.get('type')}: {edge.get('label')})")
    if not dep_edges:
        add_body("No explicit dependency links extracted.")

    add_heading_styled("9.  Wiring Diagram & Pin Configuration", level=1)
    connections = data.get("wiring_diagram", {}).get("connections", [])
    pin_headers = ["Source", "Pin", "Target", "Pin", "Protocol"]
    pin_rows = [[conn.get("source", ""), conn.get("source_pin", ""),
                  conn.get("target", ""), conn.get("target_pin", ""),
                  conn.get("protocol", "")] for conn in connections]
    add_table_from_rows(pin_headers, pin_rows)
    doc.add_page_break()

    # ===================== DATASHEETS + PAPERS =====================
    add_heading_styled("10.  Datasheets", level=1)
    ds_headers = ["Component", "Datasheet URL", "Source"]
    ds_rows = [[ds.get("component", ""), ds.get("datasheet_link", ""), ds.get("source", "")]
                for ds in data.get("datasheets", [])]
    add_table_from_rows(ds_headers, ds_rows)

    add_heading_styled("11.  Research Papers (Top 3)", level=1)
    papers = data.get("papers", [])
    rp_headers = ["Title", "Score", "Year", "URL"]
    rp_rows = [[p.get("title", "—"), f"{p.get('score', 80)}/100",
                 str(p.get("publish_year", 2022)), p.get("url", "")]
                for p in papers[:3]]
    add_table_from_rows(rp_headers, rp_rows)
    paper_summary = data.get("paper_summary", {})
    add_body(paper_summary.get("summary", "Consolidated literature compiled in research logs."))
    doc.add_page_break()

    # ===================== VALIDATION + OPTIMISATION =====================
    add_heading_styled("12.  Engineering Validation", level=1)
    for chk in validation.get("validation_checks", []):
        add_body(f"[{chk.get('status')}]  {chk.get('check')}:  {chk.get('details')}")
    for c in validation.get("contradictions", []):
        add_body(f"⚠  {c.get('conflict')} ({c.get('severity')}):  {c.get('explanation')}  "
                 f"Mitigation: {c.get('mitigation')}")

    add_heading_styled("13.  Optimisation Report", level=1)
    opt = data.get("optimization", {})
    for r in opt.get("recommendations", []):
        add_body(f"•  {r}")
    add_body(f"Optimisation score: {opt.get('optimization_score', 90)} %  |  "
             f"Cost saved: ₹{int(opt.get('saved_amount', 12)*83):,}")
    doc.add_page_break()

    # ===================== CODE =====================
    add_heading_styled("14.  Generated Controller Firmware", level=1)
    add_body("Target MCU: ESP32  |  Framework: Arduino C++", bold=True)
    code_text = generate_firmware_code(components, intent)
    for line in code_text.split("\n"):
        add_code(line)
    doc.add_page_break()

    # ===================== GANTT + CONNECTION NOTES =====================
    add_heading_styled("15.  Gantt Timeline", level=1)
    for phase in data.get("roadmap", []):
        add_body(f"Phase {phase.get('phase')}: {phase.get('title')}  "
                 f"({phase.get('duration_days')} days)  —  {phase.get('deliverable', '')}", bold=True)
        add_body(phase.get("description", ""))

    add_heading_styled("16.  Connection Assistant Notes", level=1)
    add_body("1. Ground Loop Prevention:  PCA9685 logic ground must share a common node with ESP32 GND; "
             "servo return current should route directly to the LiPo terminal block.", bold=True)
    add_body("2. I²C Pull-ups:  Place 4.7 kΩ pull-up resistors on SDA/SCL to the 3.3 V rail.", bold=True)
    add_body("3. Power Decoupling:  Install a 1000 µF low-ESR electrolytic across PCA9685 V+ terminal.", bold=True)
    doc.add_page_break()

    # ===================== AUDIT + CONCLUSION =====================
    add_heading_styled("17.  ArmorIQ Audit Trail", level=1)
    audit_trail = data.get("audit_trail", [])
    a_headers = ["Agent", "Tool", "Scope", "Status", "Receipt Hash"]
    a_rows = []
    for log in audit_trail[:14]:
        h = log.get("receipt_hash", "—")
        short = f"{h[:6]}…{h[-6:]}" if len(h) > 12 else h
        a_rows.append([log.get("agent", ""), log.get("tool", ""),
                        log.get("scope", ""), log.get("status", ""), short])
    add_table_from_rows(a_headers, a_rows)

    add_heading_styled("18.  Conclusion & Execution Readiness", level=1)
    add_body(f"Build Feasibility:  {readiness_score} %", bold=True)
    add_body(f"Cost Feasibility:  92 %", bold=True)
    add_body(f"Complexity:  {complexity}", bold=True)
    add_body("Estimated Build Time:  12 days  (4-phase roadmap)", bold=True)
    add_body("Skill Level:  Intermediate Electronics & Firmware", bold=True)
    add_body(
        f"The {project_title} project has been compiled and verified by the multi-agent system. "
        "Electrical wiring schematics, I²C logic domains, and battery discharge thresholds show full "
        "compatibility. No critical hazards were logged in the validation audit. Sourcing reflects the "
        "most cost-effective alternatives, making the project fully ready for hardware prototyping."
    )

    doc.save(file_path)
    return {"filename": filename, "url": f"/api/exports/{filename}", "status": "SUCCESS"}
